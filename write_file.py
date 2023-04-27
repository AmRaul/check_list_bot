
# https://docs-python.ru/packages/modul-python-docx-python/
import os

from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def write_files(name, show_auto):

    doc = Document()
    # Добавляем абзац
    p = doc.add_paragraph(f'{name}')
    # выравниваем текст абзаца
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # получаем объект форматирования
    fmt = p.paragraph_format
    # Форматируем:
    # добавляем отступ слева
    fmt.first_line_indent = Mm(15)
    # добавляем отступ до
    fmt.space_before = Mm(20)
    # добавляем отступ слева
    fmt.space_after = Mm(10)
    doc.add_paragraph(f'{show_auto}')
    # Перебираем папку с изображениями, добавляем каждый файл в документ и удаляем
    for filename in os.listdir('img'):
        if filename[filename.rfind(".") + 1:] in ['jpg', 'jpeg', 'png']:
            doc.add_picture(f'img/{filename}', width=Mm(150))
            os.remove(f'img/{filename}')
    doc.save(f'documents/{name}.docx')



